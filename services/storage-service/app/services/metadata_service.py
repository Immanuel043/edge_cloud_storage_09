"""
Metadata Extraction Service
Extracts metadata from various file types: images, PDFs, audio, video, documents
"""

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

import fitz  # PyMuPDF
from PIL import Image
from PIL.ExifTags import GPSTAGS, TAGS
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


class MetadataService:
    """Service for extracting metadata from files"""

    async def extract_metadata(self, file_data: bytes, mime_type: str, filename: str) -> Dict:
        """
        Extract metadata based on file type

        Args:
            file_data: File bytes
            mime_type: MIME type
            filename: Original filename

        Returns:
            Dict with metadata
        """
        try:
            if mime_type.startswith("image/"):
                return await self.extract_image_metadata(file_data)
            elif mime_type == "application/pdf":
                return await self.extract_pdf_metadata(file_data)
            elif mime_type.startswith("audio/"):
                return await self.extract_audio_metadata(file_data, filename)
            elif mime_type.startswith("video/"):
                return await self.extract_video_metadata(file_data, filename)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ]:
                return await self.extract_docx_metadata(file_data)
            else:
                return {"type": "unknown"}

        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}", exc_info=True)
            return {"error": str(e)}

    async def extract_image_metadata(self, image_data: bytes) -> Dict:
        """Extract EXIF and other metadata from images"""
        try:
            image = Image.open(io.BytesIO(image_data))

            metadata = {
                "type": "image",
                "format": image.format,
                "mode": image.mode,
                "width": image.width,
                "height": image.height,
                "size_mp": round((image.width * image.height) / 1000000, 2),
            }

            # Extract EXIF data
            exif_data = {}
            if hasattr(image, "_getexif") and image._getexif():
                exif = image._getexif()
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    try:
                        # Handle GPS data specially
                        if tag == "GPSInfo":
                            gps_data = {}
                            for gps_tag_id in value:
                                gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                                gps_data[gps_tag] = value[gps_tag_id]
                            exif_data[tag] = gps_data
                        # Convert bytes to string
                        elif isinstance(value, bytes):
                            exif_data[tag] = value.decode("utf-8", errors="ignore")
                        # Handle IFD (Image File Directory)
                        elif isinstance(value, dict):
                            exif_data[tag] = {k: str(v) for k, v in value.items()}
                        else:
                            exif_data[tag] = str(value)
                    except:
                        exif_data[tag] = str(value)

            if exif_data:
                metadata["exif"] = exif_data

                # Extract common useful fields
                if "DateTime" in exif_data:
                    metadata["date_taken"] = exif_data["DateTime"]
                if "Make" in exif_data:
                    metadata["camera_make"] = exif_data["Make"]
                if "Model" in exif_data:
                    metadata["camera_model"] = exif_data["Model"]
                if "Orientation" in exif_data:
                    metadata["orientation"] = exif_data["Orientation"]
                if "Flash" in exif_data:
                    metadata["flash"] = exif_data["Flash"]
                if "FocalLength" in exif_data:
                    metadata["focal_length"] = exif_data["FocalLength"]
                if "ISOSpeedRatings" in exif_data:
                    metadata["iso"] = exif_data["ISOSpeedRatings"]

            return metadata

        except Exception as e:
            logger.error(f"Image metadata extraction failed: {e}")
            return {"type": "image", "error": str(e)}

    async def extract_pdf_metadata(self, pdf_data: bytes) -> Dict:
        """Extract metadata from PDF files"""
        try:
            metadata = {"type": "pdf"}

            # Method 1: PyMuPDF (more comprehensive)
            try:
                doc = fitz.open(stream=pdf_data, filetype="pdf")
                pdf_meta = doc.metadata

                metadata.update(
                    {
                        "page_count": len(doc),
                        "title": pdf_meta.get("title", ""),
                        "author": pdf_meta.get("author", ""),
                        "subject": pdf_meta.get("subject", ""),
                        "keywords": pdf_meta.get("keywords", ""),
                        "creator": pdf_meta.get("creator", ""),
                        "producer": pdf_meta.get("producer", ""),
                        "creation_date": pdf_meta.get("creationDate", ""),
                        "modification_date": pdf_meta.get("modDate", ""),
                        "encrypted": doc.is_encrypted,
                        "format": pdf_meta.get("format", "PDF"),
                    }
                )

                # Get page dimensions
                if len(doc) > 0:
                    page = doc[0]
                    metadata["page_width"] = page.rect.width
                    metadata["page_height"] = page.rect.height

                doc.close()

            except Exception as e:
                logger.warning(f"PyMuPDF metadata extraction failed: {e}")

                # Fallback: PyPDF2
                try:
                    reader = PdfReader(io.BytesIO(pdf_data))
                    pdf_info = reader.metadata

                    metadata.update(
                        {
                            "page_count": len(reader.pages),
                            "title": pdf_info.get("/Title", ""),
                            "author": pdf_info.get("/Author", ""),
                            "subject": pdf_info.get("/Subject", ""),
                            "creator": pdf_info.get("/Creator", ""),
                            "producer": pdf_info.get("/Producer", ""),
                            "creation_date": pdf_info.get("/CreationDate", ""),
                            "modification_date": pdf_info.get("/ModDate", ""),
                        }
                    )
                except Exception as e2:
                    logger.warning(f"PyPDF2 metadata extraction also failed: {e2}")

            return metadata

        except Exception as e:
            logger.error(f"PDF metadata extraction failed: {e}")
            return {"type": "pdf", "error": str(e)}

    async def extract_audio_metadata(self, audio_data: bytes, filename: str) -> Dict:
        """Extract metadata from audio files"""
        try:
            # Save to temp file (mutagen needs file path)
            import tempfile

            from mutagen import File as MutagenFile

            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
                tmp.write(audio_data)
                tmp_path = tmp.name

            try:
                audio = MutagenFile(tmp_path)

                if audio is None:
                    return {"type": "audio", "error": "Unsupported audio format"}

                metadata = {
                    "type": "audio",
                    "duration": getattr(audio.info, "length", 0),
                    "bitrate": getattr(audio.info, "bitrate", 0),
                    "sample_rate": getattr(audio.info, "sample_rate", 0),
                    "channels": getattr(audio.info, "channels", 0),
                }

                # Extract tags
                if audio.tags:
                    tags = {}
                    for key, value in audio.tags.items():
                        if isinstance(value, list):
                            tags[key] = [str(v) for v in value]
                        else:
                            tags[key] = str(value)
                    metadata["tags"] = tags

                    # Extract common fields
                    title = audio.tags.get("TIT2") or audio.tags.get("title")
                    artist = audio.tags.get("TPE1") or audio.tags.get("artist")
                    album = audio.tags.get("TALB") or audio.tags.get("album")
                    genre = audio.tags.get("TCON") or audio.tags.get("genre")
                    year = audio.tags.get("TDRC") or audio.tags.get("date")

                    if title:
                        metadata["title"] = str(title[0]) if isinstance(title, list) else str(title)
                    if artist:
                        metadata["artist"] = (
                            str(artist[0]) if isinstance(artist, list) else str(artist)
                        )
                    if album:
                        metadata["album"] = str(album[0]) if isinstance(album, list) else str(album)
                    if genre:
                        metadata["genre"] = str(genre[0]) if isinstance(genre, list) else str(genre)
                    if year:
                        metadata["year"] = str(year[0]) if isinstance(year, list) else str(year)

                return metadata

            finally:
                # Clean up temp file
                import os

                try:
                    os.unlink(tmp_path)
                except:
                    pass

        except ImportError:
            logger.warning("mutagen not available for audio metadata")
            return {"type": "audio", "error": "mutagen not installed"}
        except Exception as e:
            logger.error(f"Audio metadata extraction failed: {e}")
            return {"type": "audio", "error": str(e)}

    async def extract_video_metadata(self, video_data: bytes, filename: str) -> Dict:
        """Extract metadata from video files"""
        try:
            import os
            import tempfile

            import cv2

            # Save to temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
                tmp.write(video_data)
                tmp_path = tmp.name

            try:
                cap = cv2.VideoCapture(tmp_path)

                if not cap.isOpened():
                    return {"type": "video", "error": "Failed to open video"}

                metadata = {
                    "type": "video",
                    "width": int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                    "height": int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
                    "fps": cap.get(cv2.CAP_PROP_FPS),
                    "frame_count": int(cap.get(cv2.CAP_PROP_FRAME_COUNT)),
                    "codec": int(cap.get(cv2.CAP_PROP_FOURCC)),
                }

                # Calculate duration
                if metadata["fps"] > 0:
                    metadata["duration"] = metadata["frame_count"] / metadata["fps"]

                cap.release()
                return metadata

            finally:
                try:
                    os.unlink(tmp_path)
                except:
                    pass

        except ImportError:
            logger.warning("opencv not available for video metadata")
            return {"type": "video", "error": "opencv not installed"}
        except Exception as e:
            logger.error(f"Video metadata extraction failed: {e}")
            return {"type": "video", "error": str(e)}

    async def extract_docx_metadata(self, docx_data: bytes) -> Dict:
        """Extract metadata from Word documents"""
        try:
            from docx import Document

            doc = Document(io.BytesIO(docx_data))
            core_props = doc.core_properties

            metadata = {
                "type": "document",
                "format": "docx",
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision,
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
            }

            # Count words
            word_count = 0
            for para in doc.paragraphs:
                word_count += len(para.text.split())
            metadata["word_count"] = word_count

            return metadata

        except ImportError:
            logger.warning("python-docx not available")
            return {"type": "document", "error": "python-docx not installed"}
        except Exception as e:
            logger.error(f"DOCX metadata extraction failed: {e}")
            return {"type": "document", "error": str(e)}


# Global metadata service instance
metadata_service = MetadataService()
